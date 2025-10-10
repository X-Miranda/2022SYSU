import time
import urllib.parse
import os
import pkuseg
import re
import xlrd
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import io
import pytesseract
from opencc import OpenCC
from pathlib import Path
import win32com.client
import tempfile
import pdfplumber
from tqdm import tqdm

pytesseract.pytesseract.tesseract_cmd = r'D:\ocr\tesseract.exe'

# 加载自定义词典
user_dict_path = r"user_dict.txt"
pku_seg = pkuseg.pkuseg(model_name="medicine", user_dict=user_dict_path)

# 加载停用词表
stopwords = set()
with open(r"stopwords.txt", 'r', encoding='utf-8') as f:
    for line in f:
        stopwords.add(line.strip())


def load_synonyms(synonyms_file):
    """加载同义词表，返回替换字典"""
    synonyms = {}
    with open(synonyms_file, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line or ':' not in line:
                continue
            std_word, alias_words = line.split(':', 1)
            std_word = std_word.strip()
            for word in alias_words.split(','):
                word = word.strip()
                if word:
                    synonyms[word] = std_word  # 映射到标准词
    return synonyms


# 使用示例
SYNONYMS = load_synonyms("policy_synonyms.txt")  # 替换为您的文件路径



def extract_text_from_doc(doc_path):
    """健壮的.doc文件处理器，包含多重保护机制"""
    abs_path = os.path.abspath(doc_path)
    if not os.path.exists(abs_path):
        raise FileNotFoundError(f"文件不存在: {abs_path}")

    word = None
    doc = None
    temp_dir = None
    max_retries = 2

    for attempt in range(max_retries):
        try:
            # 创建临时目录
            temp_dir = tempfile.mkdtemp()
            temp_docx = os.path.join(temp_dir, f"temp_{attempt}.docx")

            # 启动Word实例（不设置Visible属性）
            word = win32com.client.Dispatch("Word.Application")
            word.DisplayAlerts = False

            # 特殊处理文件名（避免中文/空格问题）
            safe_path = f'"{abs_path}"' if ' ' in abs_path else abs_path

            # 打开文档（显式指定文件格式）
            doc = word.Documents.Open(
                FileName=safe_path,
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False
            )

            # 保存为临时docx
            doc.SaveAs(temp_docx, FileFormat=16)
            doc.Close()

            # 提取文本
            text = extract_text_from_docx(temp_docx)
            return text

        except Exception as e:
            print(f"尝试 {attempt + 1} 失败: {str(e)}")
            if attempt == max_retries - 1:
                raise
            time.sleep(1)  # 重试前等待

        finally:
            # 确保资源释放
            if doc and hasattr(doc, 'Close'):
                try:
                    doc.Close()
                except:
                    pass
            if word and hasattr(word, 'Quit'):
                try:
                    word.Quit()
                except:
                    pass
            # 清理临时文件
            if temp_dir and os.path.exists(temp_dir):
                try:
                    for f in os.listdir(temp_dir):
                        os.remove(os.path.join(temp_dir, f))
                    os.rmdir(temp_dir)
                except:
                    pass


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    seen_text = set()  # 记录已出现的内容
    full_text = []

    # 处理段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and text not in seen_text:
            full_text.append(text)
            seen_text.add(text)

    # 处理表格（添加去重）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in seen_text:
                    full_text.append(text)
                    seen_text.add(text)

    # 处理图片中的文字（OCR）
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_bytes = rel.target_part.blob
            img = Image.open(io.BytesIO(img_bytes))
            img_text = pytesseract.image_to_string(img, lang='chi_sim')
            full_text.append(img_text)

    return "\n".join(full_text)


def extract_text_with_pdfplumber(pdf_path):
    """增强版PDF提取（自动处理文字/图片型PDF）"""
    full_text = []

    with pdfplumber.open(pdf_path, laparams={"line_overlap": 0.7}) as pdf:
        for page in pdf.pages:
            text = page.extract_text(
                x_tolerance=3,
                y_tolerance=3,
                keep_blank_chars=False,
                use_text_flow=True
            )
            if not text or len(text) < 50:
                print("启用ocr")
                img = page.to_image(resolution=300).original
                text = pytesseract.image_to_string(img, lang='chi_sim')

            if text:
                full_text.append(text)
    return "\n".join(full_text)


def word_cut(content):
    cc = OpenCC('t2s')  # 繁体转简体
    content = cc.convert(content)
    # 保留中文和AI/ai，去除其他所有英文字符
    content = re.sub(r'(?![aA][iI])[a-zA-Z]', '', content)  # 去除除AI/ai外的所有英文字母
    content = re.sub(u'[^\u4e00-\u9fa5aA]', '', content)  # 保留中文和字母A/a
    content = re.sub(r'a(?!i)|A(?!I)', '', content)  # 去除单独的a/A（后面不跟i/I的）
    words = pku_seg.cut(content)
    filtered_words = [word for word in words if word not in stopwords]
    normalized_words = [SYNONYMS.get(word, word) for word in filtered_words]
    return normalized_words

def _read_xlsx(file_path):
    """读取.xlsx文件"""
    wb = load_workbook(file_path, data_only=True)
    result = {}
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        result[sheet_name] = []
        for row in sheet.iter_rows(values_only=True):
            result[sheet_name].append([str(cell) if cell is not None else "" for cell in row])
    return result

def _read_xls(file_path):
    """读取.xls文件（旧格式）"""
    book = xlrd.open_workbook(file_path)
    result = {}
    for sheet in book.sheets():
        result[sheet.name] = []
        for row_idx in range(sheet.nrows):
            result[sheet.name].append([
                str(sheet.cell_value(row_idx, col_idx))
                for col_idx in range(sheet.ncols)
            ])
    return result

def process_attachment(attachment_path, output_folder, name_prefix):
    """
    处理单个附件文件
    :param name_prefix: 直接传入"附件：原文件名"格式
    """
    try:
        ext = os.path.splitext(attachment_path)[1].lower()
        if ext == '.docx':
            text = extract_text_from_docx(attachment_path)
        elif ext == '.doc':
            text = extract_text_from_doc(attachment_path)
        elif ext == '.pdf':
            text = extract_text_with_pdfplumber(attachment_path)
        elif ext in ('.xlsx', '.xls'):
            excel_data = _read_xlsx(attachment_path) if ext == '.xlsx' else _read_xls(attachment_path)
            text = "\n".join(["\t".join(row) for sheet in excel_data.values() for row in sheet])
        else:
            print(f"不支持的文件类型: {attachment_path}")
            return None, None

        processed_text = word_cut(text)
        output_filename = f"{name_prefix}.txt"
        output_path = os.path.join(output_folder, output_filename)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(" ".join(processed_text))
        return processed_text, output_path
    except Exception as e:
        print(f"处理附件失败 {attachment_path}: {e}")
        return None, None

def process_folder(input_folder, output_folder):
    """
    处理指定文件夹内的所有文件
    :param input_folder: 包含待处理文件的文件夹路径
    :param output_folder: 保存分词结果的文件夹路径
    """
    os.makedirs(output_folder, exist_ok=True)
    for filename in os.listdir(input_folder):
        file_path = os.path.join(input_folder, filename)
        if os.path.isfile(file_path):
            print(f"正在处理文件: {filename}")
            _, _ = process_attachment(file_path, output_folder, name_prefix=filename)
    print("文件夹内所有文件处理完成！")


# 示例：处理指定文件夹内的所有文件
input_folder = r'.\政策'  # 替换为你的输入文件夹路径
output_folder = r'分词结果'  # 输出文件夹路径
process_folder(input_folder, output_folder)