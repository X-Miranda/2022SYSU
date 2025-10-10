import time
from pathlib import Path
from selenium import webdriver
from selenium.webdriver.edge.service import Service
from selenium.webdriver.edge.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import requests
from bs4 import BeautifulSoup
import os
import pandas as pd
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Font
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from requests.exceptions import RequestException


def save_to_word(content, file_path):
    soup = BeautifulSoup(content, 'html.parser')
    doc = Document()

    # 设置默认字体样式
    default_style = doc.styles['Normal']
    default_style.font.name = 'Times New Roman'
    default_style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    for element in soup.descendants:
        if element.name == 'p':
            # 处理段落标签
            paragraph = doc.add_paragraph()
            if element.style and 'text-align: center;' in element.style:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif element.style and 'text-align: justify;' in element.style:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            if element.get('indenttext', '') == '　　':
                paragraph_format = paragraph.paragraph_format
                paragraph_format.first_line_indent = Inches(0.25)  # 设置首行缩进为0.25英寸

            text = element.get_text(strip=True)
            if element.strong:
                run = paragraph.add_run(text)
                run.bold = True
            else:
                paragraph.add_run(text)

        elif element.name == 'h1' or element.name == 'h2':
            # 处理标题标签
            heading_level = 0 if element.name == 'h1' else 1
            doc.add_heading(element.get_text(strip=True), level=heading_level)
        elif element.name == 'img':
            # 处理图片标签
            img_url = element['src']
            img_request = requests.get(img_url)
            img_request.raise_for_status()
            img_name = os.path.basename(img_url)
            img_path = os.path.join('./图片', img_name)
            with open(img_path, 'wb') as img_file:
                img_file.write(img_request.content)
            doc.add_picture(img_path)

    # 检查Word文档是否存在，如果存在则添加序号
    base, extension = os.path.splitext(file_path)
    counter = 1
    while os.path.exists(file_path):
        file_path = f"{base}_{counter}{extension}"
        counter += 1
    doc.save(file_path)

def save_file(response, file_path):
    # 检查文件是否存在，如果存在则添加序号
    base, extension = os.path.splitext(file_path)
    counter = 1
    new_file_path = file_path
    while Path(new_file_path).exists():
        new_file_path = f"{base}_{counter}{extension}"
        counter += 1
    # 保存文件
    with open(new_file_path, 'wb') as f:
        f.write(response.content)
    return new_file_path

def get_policy_details(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    details = {
        "政策标题": "",
        "印发单位": "",
        "发布日期": "",
        "文号": "",
        "政策正文": "",
        "分词结果文件": "",
        "附件1": ""
    }
    try:
        # 提取政策标题
        for t_tag in soup.find_all('p'):
            if '标题' in t_tag.text:
                title = t_tag.find('span').text.strip()
                details["政策标题"] = title
                break
    except Exception:
        print("未找到政策标题")

    try:
        # 提取印发单位
        for p_tag in soup.find_all('p'):
            if '印发单位' in p_tag.text:
                issue_unit = p_tag.find('span').text.strip()
                details["印发单位"] = issue_unit
                break
    except Exception:
        print("未找到印发单位")

    try:
        # 提取发布日期
        for d_tag in soup.find_all('p'):
            if '发布日期' in d_tag.text:
                publish_date = d_tag.find('span').text.strip()
                details["发布日期"] = publish_date
                break
    except Exception:
        print("未找到发布日期")

    try:
        # 提取文号
        doc_number_tag = soup.find('p', class_='wh')
        doc_number = doc_number_tag.find('span').text.strip()
        details["文号"] = doc_number
    except Exception:
        print("未找到文号")

    try:
        # 爬取政策内容
        content_div = soup.find(class_='article_con')
        if content_div:
            content_html = content_div.prettify()
            # 提取内容标题
            content_title_tag = soup.find('h1', class_='article_t')
            content_title = content_title_tag.text.strip()
            word_file_name = f"{content_title}.docx"
            word_file_path = os.path.join('./政策', word_file_name)
            save_to_word(content_html, word_file_path)
            details["政策正文"] = word_file_name
    except Exception:
        print("内容未找到")

    try:
        # 爬取附件
        download_link_list = soup.find_all('a', class_='nfw-cms-attachment')
        index = 1
        for download_link in download_link_list:
            download_url = download_link.get('href')
            file_name = download_link.get('alt')
            if not file_name:
                file_name = f"附件{index}.pdf"  # 如果没有alt属性，使用默认文件名
            file_path = os.path.join('./附件', file_name)
            try:
                response = requests.get(download_url)
                response.raise_for_status()
                # 保存文件并获取可能已修改的文件名
                file_path = save_file(response, file_path)
                print(f"文件下载成功，保存为：{file_path}")
                details[f"附件{index}"] = os.path.basename(file_path)  # 使用basename获取文件名
                index += 1
            except requests.RequestException as e:
                print(f"文件下载失败: {e}")
    except Exception:
        print("附件未找到")

    return details

'''def read_existing_policies(file_path):
    df_read = pd.read_excel(file_path)
    if df_read.empty:
        return set()
    existing_policies = set(df_read['政策标题'].tolist())
    return existing_policies'''

if __name__ == '__main__':
    # 创建 Service 对象，指定驱动程序的路径
    service = Service(executable_path='C:/Program Files (x86)/Microsoft/Edge/Application/msedgedriver.exe')

    edge_options = Options()
    # 1. 开启开发者模式
    edge_options.add_experimental_option('excludeSwitches', ['enable-automation'])
    # 2. 禁用启用Blink运行时的功能
    edge_options.add_argument('--disable-blink-features=AutomationControlled')
    # 设置保持浏览器不自动关闭
    edge_options.add_experimental_option("detach", True)

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36 Edg/130.0.0.0'
    }

    # 使用 Service 对象创建 WebDriver 实例
    driver = webdriver.Edge(service=service)
    driver.get('https://www.cnbayarea.org.cn/policy/policy%20release/policies/')
    driver.maximize_window()


    driver.find_element(By.XPATH, "//*[@id='geographical_label']/span[1]").click()
    driver.find_element(By.XPATH, "//*[@id='policy_label']/span[1]").click()
    time.sleep(1)

    policies_data = []
    file_list = []
    # existing_policies = read_existing_policies('policies_data.xlsx')
    i = 1

    while True:
        policy_list = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, 'policy_list')))
        # 获取所有的政策链接
        link_list = policy_list.find_elements(By.TAG_NAME, 'a')
        # 遍历链接并抓取数据
        for link in link_list:
            url = link.get_attribute('href')
            print(url)
            response = requests.get(url, headers=headers)
            if response.status_code == 200:
                # if title in existing_policies:
                #    continue
                # 使用BeautifulSoup解析政策页面内容
                policies_data.append(get_policy_details(response.text))
        print(f"page:{i}")
        print(len(policies_data))


        if not driver.find_elements(By.XPATH, "//a[@class='next']"):
            break  # 如果没有下一页按钮，结束循环

        next_page_button = driver.find_element(By.XPATH, "//a[@class='next']")
        # 检查是否还有下一页

        next_page_button.click()
        i += 1
        time.sleep(2)  # 等待3秒，等待页面加载

    df = pd.DataFrame(policies_data)
    df.to_excel("policies_data.xlsx", index=False)

    wb = load_workbook("policies_data.xlsx")
    ws = wb.active
    base_path = './附件/'
    content_path = './政策/'

    for row, data in enumerate(policies_data, start=2):
        file_content = data.get('政策正文', '')

        # 为每个附件设置超链接
        attachment_index = 1
        while f"附件{attachment_index}" in data:
            file_name = data.get(f"附件{attachment_index}", '')
            if file_name:  # 如果存在附件
                dr_file_path = os.path.join(base_path, file_name)  # 构建完整的文件路径
                column_letter = chr(ord('G') + attachment_index - 1)  # 计算列号
                ws[f"{column_letter}{row}"].value = file_name  # 设置单元格显示的文本为文件名
                ws[f"{column_letter}{row}"].hyperlink = dr_file_path  # 设置超链接
                ws[f"{column_letter}{row}"].font = Font(color="0000FF", underline="single")  # 设置超链接颜色为蓝色
            attachment_index += 1

        if file_content:
            content_file_path = os.path.join(content_path, file_content)
            ws[f"E{row}"].hyperlink = content_file_path  # 设置超链接
            ws[f"E{row}"].font = Font(color="009900")  # 设置超链接颜色为绿色
            ws[f"E{row}"].value = file_content  # 设置单元格显示的文本为文件名

    wb.save("policies_data.xlsx")
    wb.close()

    print(f"数据抓取完成，已保存到 policies_data.xlsx")